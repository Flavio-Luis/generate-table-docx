from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from num2words import num2words
import os
from copy import deepcopy 

default_columns = ["Número", "Unid. Med", "Descrição", "Marca", "Modelo", "Valor Unit", "Valor Descrito", "Qtde Unit", "Total"]

def set_margins(doc, top=1, bottom=1, left=1, right=1):
    section = doc.sections[0]
    section.top_margin = Pt(top * 28.3465) 
    section.bottom_margin = Pt(bottom * 28.3465)
    section.left_margin = Pt(left * 28.3465)
    section.right_margin = Pt(right * 28.3465)

def generate_table(products, file_name, library, extra_column=None):
    doc = Document()
    set_margins(doc, top=1, bottom=1, left=1, right=1)

    columns = deepcopy(default_columns)
    if extra_column:
        columns.extend(extra_column)

    table = doc.add_table(rows=len(products) + 1, cols=len(columns), style="Table Grid")

    header_table = table.rows[0].cells
    for i, name_columns in enumerate(columns):
        header_table[i].text = name_columns

    total_value = 0
    for i, product in enumerate(products):
        row_cell = table.rows[i + 1].cells
        for j, column in enumerate(columns):
            value = product.get(column, '')
            row_cell[j].text = str(value)
            if column == "Valor Unit":
                valor_unit = product.get(column, 0)
            if column == "Qtde Unit":
                qtde_unit = product.get(column, 0)
            if column == "Total":
                value = float(valor_unit) * float(qtde_unit)
                row_cell[j].text = str(value)
                total_value += float(value)

    total_row = table.add_row().cells
    total_row[0].merge(total_row[-1])

    reais = int(total_value)
    centavos = int(round((total_value - reais) * 100))
    
    real_extensive = num2words(reais, lang="pt_BR", to='currency')
    centavos_extensive = num2words(centavos, lang="pt_BR", to='currency')

    join_extensive = f"{real_extensive} e {centavos_extensive}"

    total_row[0].text = f"TOTAL: {join_extensive}   |   R${total_value:.2f}"

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:

                paragraph.paragraph_format.left_indent = Pt(1)
                paragraph.paragraph_format.right_indent = Pt(1)

                paragraph.paragraph_format.space_before = Pt(1)
                paragraph.paragraph_format.space_after = Pt(1)
                paragraph.paragraph_format.line_spacing = Pt(12)

    file_path = os.path.join(library, file_name)
    doc.save(file_path)

products = [
    {"Número": 1, "Unid. Med": "UN", "Descrição": "PC GAMER", "Marca": "Ryzhen", "Modelo": "AMD", "Valor Unit": 35.95, "Valor Descrito": "teste", "Qtde Unit": 1, "Total": 0},
    {"Número": 2, "Unid. Med": "UN", "Descrição": "Calculadora", "Marca": "Ryzhen", "Modelo": "AMD", "Valor Unit": 100.00, "Valor Descrito": "teste", "Qtde Unit": 75, "Total": 0}
]

generate_table(products, 'tabela_produtos.docx', library="C:/GitHub/Python/gerar_word_tables")
